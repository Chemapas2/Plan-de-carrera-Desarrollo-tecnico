import io
import re
import tempfile
from html import escape
from pathlib import Path
from urllib.parse import quote_plus, urlparse

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer

st.set_page_config(
    page_title="Assessment + Plan de Carrera Técnico",
    page_icon="📈",
    layout="wide",
)

QUESTION_GROUPS = {
    "Potencial y proyección": [
        "¿Cuál es el potencial real del evaluado para asumir funciones técnicas de mayor complejidad?",
        "¿En qué áreas muy específicas o exigentes podría ser formado con mayor probabilidad de éxito?",
        "¿Qué combinación de fortalezas actuales le da más proyección a medio plazo?",
        "¿Qué evidencias indican que puede crecer por encima del estándar actual del equipo?",
        "¿Qué nivel de ambición técnica parece realista en su siguiente etapa de desarrollo?",
    ],
    "Fortalezas y especialización": [
        "Además de las áreas ya propuestas, ¿qué otras áreas tienen mayor probabilidad de éxito para especialización?",
        "¿Qué fortalezas diferenciales conviene consolidar antes de ampliar su ámbito de especialización?",
        "¿En qué tronco o conjunto de indicadores destaca de forma más consistente?",
        "¿Qué fortalezas se pueden convertir en una línea de excelencia técnica?",
        "¿Qué áreas merecen un plan de especialización avanzada y no solo de mantenimiento?",
    ],
    "Brechas y riesgos": [
        "¿Cuáles son las brechas que más limitan hoy su desarrollo profesional?",
        "¿Qué debilidades tienen mayor impacto operativo en su desempeño actual?",
        "¿Qué área débil conviene corregir primero para evitar efecto arrastre sobre otras competencias?",
        "¿Dónde aparece el mayor gap frente a referencia y frente al benchmark del equipo?",
        "¿Qué riesgos tendría forzar una especialización sin cerrar antes ciertas bases?",
    ],
    "Plan de acción y formación": [
        "¿Qué plan de desarrollo de 6 meses tendría más sentido para este perfil?",
        "¿Qué plan de desarrollo de 12 meses tendría más sentido para este perfil?",
        "¿Qué formación base debería completar antes de abordar contenidos avanzados?",
        "¿Qué combinación de formación interna y externa parece más eficiente para este caso?",
        "¿Qué indicadores deberían revisarse de nuevo en la próxima evaluación para confirmar progreso?",
    ],
}
QUESTION_TO_GROUP = {q: g for g, qs in QUESTION_GROUPS.items() for q in qs}

# Clasificación principal: se usa la lógica del Excel (ENLACES DATOS!N44),
# con cortes leídos de ENLACES DATOS!M56:M61 y ratio ENLACES DATOS!K37.
PROFILE_THRESHOLDS = {
    "Basico": 0.4,
    "Controla": 0.5,
    "Supera": 0.6,
    "Certificado": 0.7,
    "Excelente": 0.8,
    "Master": 0.9,
}

AREA_KEYWORDS = {
    "Nutrición": ["nutricion", "nutrición", "producto", "aditivo", "agua", "alimentacion", "alimentación", "cebo", "racion", "ración", "pienso", "forraje", "ruminal", "formulacion", "formulación"],
    "Patología": ["patologia", "patología", "metabolica", "metabólica", "infecciosa", "parasitaria", "antibioterapia", "diagnost", "inmunidad", "salud", "microbioma", "sanidad"],
    "Manejo": ["bioseguridad", "instalacion", "instalación", "ventilacion", "ventilación", "arranque", "produccion", "producción", "reposicion", "reposición", "manejo", "granja", "housing", "ordeño", "ordeno", "reproduccion", "reproducción", "bienestar"],
    "Herramientas": ["datos", "estadistica", "estadística", "crm", "power bi", "bbdd", "digitalizacion", "digitalización", "informe", "ingles", "inglés", "software", "programa", "kpi", "analytics", "automat"],
}

DEFAULT_CATALOG_PATTERNS = [
    "*Formacion*.xlsx",
    "*formacion*.xlsx",
    "*Especializacion*.xlsx",
    "*especializacion*.xlsx",
    "*Resumen*.xlsx",
    "*resumen*.xlsx",
]
SEARCH_HEADERS = {"User-Agent": "Mozilla/5.0"}


def safe_float(value):
    if value in (None, ""):
        return None
    try:
        return float(value)
    except Exception:
        return None


def pct(value):
    value = safe_float(value)
    if value is None:
        return None
    return round(value * 100.0, 1)


def format_pct(value):
    p = pct(value)
    return "No disponible" if p is None else f"{p:.1f}%"



def canonical_trunk_name(name):
    n = normalize_text(name)
    if "aliment" in n or "nutric" in n:
        return "Alimentación"
    if "sanidad" in n or "patolog" in n:
        return "Sanidad"
    if "manejo" in n:
        return "Manejo"
    if "herramient" in n:
        return "Herramientas"
    return str(name) if name not in (None, "") else "General"


def safe_indicator_reference_raw(row):
    weight = safe_float(row.get("weight")) or 0.0
    objective_weighted = safe_float(row.get("objective_weighted"))
    if weight > 0 and objective_weighted is not None:
        return objective_weighted / weight
    return None


def safe_indicator_benchmark_raw(row):
    score_raw = safe_float(row.get("score_raw"))
    vs_bbdd = safe_float(row.get("vs_bbdd"))
    if score_raw is None or vs_bbdd in (None, 0):
        return None
    return score_raw / vs_bbdd


def build_global_radar_figure(trunks_df):
    radar_df = trunks_df.copy()
    radar_df["tronco_display"] = radar_df["tronco"].apply(canonical_trunk_name)
    radar_df["Técnico"] = radar_df["score_raw_avg"].round(2)
    radar_df["Referencia"] = 3.0
    radar_df["Benchmark aprox."] = radar_df["benchmark_raw_avg"].fillna(0).round(2)

    fig = go.Figure()
    for series in ["Técnico", "Referencia", "Benchmark aprox."]:
        fig.add_trace(
            go.Scatterpolar(
                r=radar_df[series].tolist(),
                theta=radar_df["tronco_display"].tolist(),
                fill="toself",
                name=series,
            )
        )
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 6])),
        showlegend=True,
        height=420,
        margin=dict(l=30, r=30, t=50, b=30),
        title="Radar global por áreas",
    )
    return fig


def build_area_radar_figure(indicators_df, area_name):
    canon = indicators_df["tronco"].apply(canonical_trunk_name)
    area_df = indicators_df[canon == area_name].copy()
    if area_df.empty:
        return None

    area_df["Referencia raw"] = area_df.apply(safe_indicator_reference_raw, axis=1)
    area_df["Benchmark raw"] = area_df.apply(safe_indicator_benchmark_raw, axis=1)
    area_df["Técnico"] = area_df["score_raw"].fillna(0).round(2)
    area_df["Referencia"] = area_df["Referencia raw"].fillna(0).round(2)
    area_df["Benchmark"] = area_df["Benchmark raw"].fillna(0).round(2)

    fig = go.Figure()
    for series in ["Técnico", "Referencia", "Benchmark"]:
        fig.add_trace(
            go.Scatterpolar(
                r=area_df[series].tolist(),
                theta=area_df["indicator"].tolist(),
                fill="toself",
                name=series,
            )
        )
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 6])),
        showlegend=False,
        height=420,
        margin=dict(l=20, r=20, t=50, b=20),
        title=f"Radar de {area_name}",
    )
    return fig


def figure_to_png_bytes(fig):
    try:
        return fig.to_image(format="png", width=1100, height=850, scale=2)
    except Exception:
        return None


def build_report_chart_pack(analysis):
    charts = []
    global_fig = build_global_radar_figure(analysis["trunks"])
    charts.append(("Radar global por áreas", global_fig))
    for area in ["Alimentación", "Sanidad", "Manejo", "Herramientas"]:
        fig = build_area_radar_figure(analysis["indicators"], area)
        if fig is not None:
            charts.append((f"Radar de {area}", fig))
    return charts


def normalize_text(text):
    text = (text or "").strip().lower()
    replacements = {"á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u", "ü": "u", "ñ": "n"}
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return re.sub(r"\s+", " ", text)


def infer_area_from_text(text):
    text_n = normalize_text(text)
    scores = {area: sum(1 for kw in kws if kw in text_n) for area, kws in AREA_KEYWORDS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "General"


def infer_species_subspecies_from_name(name):
    n = normalize_text(Path(str(name)).stem.replace("_", " ").replace("-", " "))
    species, subspecies = "General", "General"
    if "avicultura" in n:
        species = "Avicultura"
        if "puesta" in n:
            subspecies = "Puesta"
        elif "broiler" in n or "carne" in n:
            subspecies = "Carne"
    elif "porcino" in n:
        species = "Porcino"
        if "industrial" in n:
            subspecies = "Industrial"
    elif "vacuno" in n and ("lechero" in n or "leche" in n):
        species = "Vacuno Lechero"
        if "intensivo" in n:
            subspecies = "Intensivo"
    elif "vacuno" in n and "carne" in n:
        species = "Vacuno Carne"
        if "intensivo" in n:
            subspecies = "Intensivo"
    elif "ovino" in n or "caprino" in n:
        species = "Ovino-Caprino"
    elif "rumiante" in n:
        species = "Rumiantes"
    return species, subspecies


def species_match(selected_species, row_species):
    s = normalize_text(selected_species)
    r = normalize_text(row_species)
    if not s or not r:
        return False
    return s in r or r in s


def area_match(target_area, row_area, row_text=""):
    t = normalize_text(target_area)
    r = normalize_text(row_area)
    if t == r:
        return True
    text = normalize_text(row_text)
    return any(kw in text for kw in AREA_KEYWORDS.get(target_area, []))


def detect_header_row(ws, max_scan=10):
    for r in range(1, min(ws.max_row, max_scan) + 1):
        vals = [normalize_text(ws.cell(r, c).value) for c in range(1, min(ws.max_column, 14) + 1)]
        joined = " | ".join([v for v in vals if v])
        if "nombre del programa" in joined or "especializacion" in joined or "especialización" in joined:
            return r, vals
    return None, []


def map_headers(headers):
    mapping = {}
    for idx, header in enumerate(headers, start=1):
        h = normalize_text(header)
        if h == "especie":
            mapping["species"] = idx
        elif "subespecie" in h:
            mapping["subspecies"] = idx
        elif "tipo de formacion" in h or "tipo de formación" in h:
            mapping["type"] = idx
        elif "nombre del programa" in h or "especializacion" in h or "especialización" in h:
            mapping["program"] = idx
        elif "duracion" in h or "duración" in h:
            mapping["duration"] = idx
        elif "modalidad" in h:
            mapping["modality"] = idx
        elif "ubicacion" in h or "ubicación" in h:
            mapping["location"] = idx
        elif "institucion" in h or "institución" in h or "centros de formacion" in h or "centros de formación" in h:
            mapping["institution"] = idx
        elif "contenido" in h:
            mapping["content"] = idx
        elif "enlace" in h or "contacto" in h or "web" in h or "url" in h:
            mapping["link"] = idx
    return mapping


def parse_catalog_sheet(ws, source_name):
    header_row, headers = detect_header_row(ws)
    if not header_row:
        return []
    mapping = map_headers(headers)
    if "program" not in mapping:
        return []

    default_species, default_subspecies = infer_species_subspecies_from_name(source_name)
    rows = []
    current_species = default_species

    for r in range(header_row + 1, ws.max_row + 1):
        values = {key: ws.cell(r, col).value for key, col in mapping.items()}
        if all(v in (None, "") for v in values.values()):
            continue

        if values.get("species") not in (None, ""):
            current_species = str(values["species"]).strip()

        program = values.get("program")
        if program in (None, ""):
            continue

        row_species = current_species if "species" in mapping else default_species
        row_subspecies = str(values.get("subspecies")).strip() if values.get("subspecies") not in (None, "") else default_subspecies
        content = str(values.get("content")).strip() if values.get("content") not in (None, "") else ""
        institution = str(values.get("institution")).strip() if values.get("institution") not in (None, "") else "No disponible"
        row_type = str(values.get("type")).strip() if values.get("type") not in (None, "") else "Curso"
        link = str(values.get("link")).strip() if values.get("link") not in (None, "") else "No disponible"
        program_text = str(program).strip()

        area = infer_area_from_text(" ".join([program_text, content, institution, str(row_species)]))
        rows.append({
            "species": row_species,
            "subspecies": row_subspecies,
            "type": row_type,
            "program": program_text,
            "duration": str(values.get("duration")).strip() if values.get("duration") not in (None, "") else "No disponible",
            "modality": str(values.get("modality")).strip() if values.get("modality") not in (None, "") else "No disponible",
            "location": str(values.get("location")).strip() if values.get("location") not in (None, "") else "No disponible",
            "institution": institution,
            "content": content,
            "link": link,
            "area": area,
            "source_file": source_name,
        })
    return rows


@st.cache_data(show_spinner=False)
def load_catalog_from_path(path_str):
    path = Path(path_str)
    wb = load_workbook(path, data_only=True)
    rows = []
    for sheet_name in wb.sheetnames:
        rows.extend(parse_catalog_sheet(wb[sheet_name], path.name))
    return pd.DataFrame(rows)


def parse_uploaded_catalog(uploaded_file):
    suffix = Path(uploaded_file.name).suffix or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name
    wb = load_workbook(tmp_path, data_only=True)
    rows = []
    for sheet_name in wb.sheetnames:
        rows.extend(parse_catalog_sheet(wb[sheet_name], uploaded_file.name))
    return pd.DataFrame(rows)


def autodiscover_catalogs(exclude_names=None):
    exclude_names = set(exclude_names or [])
    paths = []
    for pattern in DEFAULT_CATALOG_PATTERNS:
        paths.extend(Path(".").glob(pattern))
        paths.extend(Path("/mnt/data").glob(pattern))
    unique = []
    seen = set()
    for p in paths:
        if p.name in exclude_names:
            continue
        if p.name not in seen:
            seen.add(p.name)
            unique.append(p)
    return unique


def build_indicator_frame(ref_ws, eval_ws):
    rows = []
    current_area = None
    for ref_row, eval_row in zip(range(4, 29), range(9, 34)):
        area = ref_ws[f"B{ref_row}"].value
        if area:
            current_area = str(area).strip()
        indicator = ref_ws[f"C{ref_row}"].value
        if not indicator:
            continue
        indicator = str(indicator).strip()
        weight = safe_float(ref_ws[f"D{ref_row}"].value) or 0.0
        objective_raw = safe_float(ref_ws[f"E{ref_row}"].value)
        objective_weighted = safe_float(ref_ws[f"F{ref_row}"].value)
        max_weighted = safe_float(ref_ws[f"G{ref_row}"].value)
        bbdd_raw = safe_float(ref_ws[f"H{ref_row}"].value)
        bbdd_weighted = safe_float(ref_ws[f"I{ref_row}"].value)
        raw_score = safe_float(eval_ws[f"D{eval_row}"].value)
        weighted_score = None if raw_score is None else raw_score * weight
        if objective_weighted is None and objective_raw is not None:
            objective_weighted = objective_raw * weight
        if max_weighted is None:
            max_weighted = 4 * weight
        if bbdd_weighted is None and bbdd_raw is not None:
            bbdd_weighted = bbdd_raw * weight
        rows.append({
            "tronco": current_area,
            "indicator": indicator,
            "weight": weight,
            "score_raw": raw_score,
            "score_weighted": weighted_score,
            "objective_weighted": objective_weighted,
            "max_weighted": max_weighted,
            "bbdd_weighted": bbdd_weighted,
            "vs_goal": None if not objective_weighted or weighted_score is None else weighted_score / objective_weighted,
            "vs_max": None if not max_weighted or weighted_score is None else weighted_score / max_weighted,
            "vs_bbdd": None if not bbdd_weighted or weighted_score is None else weighted_score / bbdd_weighted,
        })
    return pd.DataFrame(rows)


def summarise_trunks(indicators_df):
    rows = []
    for trunk, grp in indicators_df.groupby("tronco", dropna=False):
        score_weighted = grp["score_weighted"].sum(skipna=True)
        objective_weighted = grp["objective_weighted"].sum(skipna=True)
        max_weighted = grp["max_weighted"].sum(skipna=True)
        bbdd_weighted = grp["bbdd_weighted"].sum(skipna=True)
        benchmark_avg = None
        tmp = (grp["score_raw"] / grp["vs_bbdd"]).replace([float("inf"), -float("inf")], pd.NA).dropna()
        if not tmp.empty:
            benchmark_avg = tmp.mean()
        rows.append({
            "tronco": trunk,
            "score_raw_avg": grp["score_raw"].mean(skipna=True),
            "benchmark_raw_avg": benchmark_avg,
            "vs_goal": None if not objective_weighted else score_weighted / objective_weighted,
            "vs_max": None if not max_weighted else score_weighted / max_weighted,
            "vs_bbdd": None if not bbdd_weighted else score_weighted / bbdd_weighted,
        })
    return pd.DataFrame(rows)


def summarise_global(indicators_df):
    score_weighted = indicators_df["score_weighted"].sum(skipna=True)
    objective_weighted = indicators_df["objective_weighted"].sum(skipna=True)
    max_weighted = indicators_df["max_weighted"].sum(skipna=True)
    bbdd_weighted = indicators_df["bbdd_weighted"].sum(skipna=True)
    return {
        "score_raw_avg": indicators_df["score_raw"].mean(skipna=True),
        "vs_goal": None if not objective_weighted else score_weighted / objective_weighted,
        "vs_max": None if not max_weighted else score_weighted / max_weighted,
        "vs_bbdd": None if not bbdd_weighted else score_weighted / bbdd_weighted,
    }


def get_excel_profile_thresholds(wb):
    thresholds = []
    ws = wb["ENLACES DATOS"] if "ENLACES DATOS" in wb.sheetnames else None
    if ws is not None:
        for r in range(56, 62):
            rank_value = safe_float(ws[f"O{r}"].value)
            label = ws[f"P{r}"].value
            cutoff = safe_float(ws[f"M{r}"].value)
            if rank_value is None or label in (None, "") or cutoff is None:
                continue
            thresholds.append({
                "rank": int(rank_value),
                "label": str(label).strip(),
                "cutoff": cutoff,
            })
    if thresholds:
        return sorted(thresholds, key=lambda x: x["cutoff"], reverse=True)
    return [
        {"rank": 6, "label": "Master", "cutoff": 0.9},
        {"rank": 5, "label": "Excelente", "cutoff": 0.8},
        {"rank": 4, "label": "Certificado", "cutoff": 0.7},
        {"rank": 3, "label": "Supera", "cutoff": 0.6},
        {"rank": 2, "label": "Controla", "cutoff": 0.5},
        {"rank": 1, "label": "Basico", "cutoff": 0.4},
    ]


def classify_profile_from_ratio(ratio_value, thresholds):
    ratio_value = safe_float(ratio_value)
    if ratio_value is None or pd.isna(ratio_value):
        return {"label": "No disponible", "rank": None, "ratio": None}
    rank_value = sum(1 for item in thresholds if ratio_value > item["cutoff"])
    label = "No disponible"
    for item in thresholds:
        if int(item["rank"]) == int(rank_value):
            label = item["label"]
            break
    return {"label": label, "rank": int(rank_value), "ratio": float(ratio_value)}


def extract_excel_profile_info(wb, global_summary):
    thresholds = get_excel_profile_thresholds(wb)
    ratio_value = None
    rank_value = None
    label_value = None

    if "ENLACES DATOS" in wb.sheetnames:
        ws = wb["ENLACES DATOS"]
        ratio_value = safe_float(ws["K37"].value)
        rank_value = safe_float(ws["M44"].value)
        label_value = ws["N44"].value

    if ratio_value is None:
        ratio_value = safe_float(global_summary.get("vs_max"))

    if rank_value is None and ratio_value is not None:
        rank_value = sum(1 for item in thresholds if ratio_value > item["cutoff"])

    if label_value in (None, "") and rank_value is not None:
        for item in thresholds:
            if int(item["rank"]) == int(rank_value):
                label_value = item["label"]
                break

    return {
        "profile_ratio": ratio_value,
        "profile_rank": int(rank_value) if rank_value is not None else None,
        "profile_label": str(label_value).strip() if label_value not in (None, "") else "No disponible",
        "profile_thresholds": thresholds,
    }

def classify_profile(avg_raw):
    if avg_raw is None or pd.isna(avg_raw):
        return "No disponible"
    if avg_raw < PROFILE_THRESHOLDS["Basico"]:
        return "Basico"
    if avg_raw < PROFILE_THRESHOLDS["Controla"]:
        return "Controla"
    if avg_raw < PROFILE_THRESHOLDS["Supera"]:
        return "Supera"
    if avg_raw < PROFILE_THRESHOLDS["Certificado"]:
        return "Certificado"
    if avg_raw < PROFILE_THRESHOLDS["Excelente"]:
        return "Excelente"
    return "Master"

def parse_assessment(file):
    suffix = Path(file.name).suffix or ".xlsm"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file.getvalue())
        temp_path = tmp.name
    wb = load_workbook(temp_path, data_only=True, keep_vba=True)
    if "REFERENCIAS" not in wb.sheetnames or "EVALUACION" not in wb.sheetnames:
        raise ValueError("El archivo no contiene las hojas 'REFERENCIAS' y 'EVALUACION' esperadas.")
    indicators_df = build_indicator_frame(wb["REFERENCIAS"], wb["EVALUACION"])
    if indicators_df.empty:
        raise ValueError("No se han podido leer los indicadores del assessment.")
    if indicators_df["score_raw"].notna().sum() < 20:
        raise ValueError("El assessment no tiene suficientes puntuaciones válidas. Abre y guarda el Excel en origen y vuelve a subirlo.")
    invalid = indicators_df[indicators_df["score_raw"].notna() & ((indicators_df["score_raw"] < 0) | (indicators_df["score_raw"] > 6))]
    if not invalid.empty:
        raise ValueError("Hay puntuaciones fuera del rango 0–6.")
    eval_ws = wb["EVALUACION"]
    global_summary = summarise_global(indicators_df)
    profile_info = extract_excel_profile_info(wb, global_summary)
    return {
        "name": str(eval_ws["C2"].value or file.name).strip(),
        "species_from_file": str(eval_ws["C4"].value or "No disponible").strip(),
        "date": str(eval_ws["C6"].value or "No disponible"),
        "indicators": indicators_df,
        "trunks": summarise_trunks(indicators_df),
        "global": global_summary,
        "profile": profile_info["profile_label"],
        "profile_rank": profile_info["profile_rank"],
        "profile_ratio": profile_info["profile_ratio"],
        "profile_thresholds": profile_info["profile_thresholds"],
    }

def rank_strengths(indicators_df, top_n=5):
    tmp = indicators_df.copy()
    tmp["strength_score"] = tmp["vs_goal"].fillna(0) * 0.45 + tmp["vs_bbdd"].fillna(0) * 0.35 + tmp["vs_max"].fillna(0) * 0.20
    return tmp.sort_values("strength_score", ascending=False).head(top_n)


def rank_weaknesses(indicators_df, top_n=5):
    tmp = indicators_df.copy()
    tmp["gap_score"] = (
        (1 - tmp["vs_goal"].fillna(0)).clip(lower=0) * 0.55
        + (1 - tmp["vs_bbdd"].fillna(0)).clip(lower=0) * 0.35
        + ((3 - tmp["score_raw"].fillna(0)) / 3).clip(lower=0) * 0.10
    )
    return tmp.sort_values("gap_score", ascending=False).head(top_n)


def build_priority_areas(indicators_df, top_n=4):
    weak = rank_weaknesses(indicators_df, top_n=8).copy()
    weak["priority"] = weak["gap_score"]
    return weak.groupby("tronco", as_index=False)["priority"].mean().sort_values("priority", ascending=False).head(top_n)


def build_specialization_areas(indicators_df, top_n=3):
    strong = rank_strengths(indicators_df, top_n=8).copy()
    strong["priority"] = strong["strength_score"]
    return strong.groupby("tronco", as_index=False)["priority"].mean().sort_values("priority", ascending=False).head(top_n)


def recommend_internal_courses(catalog_df, species, subspecies, area, max_items=4):
    if catalog_df.empty:
        return pd.DataFrame()
    cat = catalog_df.copy()
    cat["match_score"] = 0.0
    cat.loc[cat["species"].apply(lambda x: species_match(species, x)), "match_score"] += 3
    if subspecies and normalize_text(subspecies) != "general":
        cat.loc[cat["subspecies"].apply(lambda x: species_match(subspecies, x) or normalize_text(x) == "general"), "match_score"] += 1
    cat.loc[
        cat.apply(lambda r: area_match(area, r["area"], " ".join([str(r["program"]), str(r["content"]), str(r["institution"])])), axis=1),
        "match_score",
    ] += 3
    cat.loc[cat["type"].astype(str).str.contains("master|máster|postgrado|diplom", case=False, na=False), "match_score"] += 0.2
    filtered = (
        cat[cat["match_score"] > 0]
        .sort_values(["match_score", "program"], ascending=[False, True])
        .drop_duplicates(subset=["program", "institution"])
        .head(max_items)
    )
    return filtered[["program", "institution", "duration", "modality", "location", "link", "area", "source_file", "type"]]


@st.cache_data(show_spinner=False, ttl=21600)
def web_search_courses(species, area, subspecies, max_items=3):
    area_query = {
        "Nutrición": "nutrition animal science feed formulation course master",
        "Patología": "animal health pathology biosecurity veterinary course diploma",
        "Manejo": "livestock management welfare facilities production course",
        "Herramientas": "data analytics digitalization livestock power bi course",
        "General": "animal science veterinary production course",
    }.get(area, "animal science veterinary production course")
    queries = [
        f'"{species}" {subspecies} {area} curso master formacion veterinaria',
        f'"{species}" {area} {area_query}',
        f'"{species}" "{area}" training course university',
    ]

    results = []
    seen = set()
    for query in queries:
        try:
            url = f"https://html.duckduckgo.com/html/?q={quote_plus(query)}"
            resp = requests.get(url, headers=SEARCH_HEADERS, timeout=15)
            if resp.status_code != 200:
                continue
            soup = BeautifulSoup(resp.text, "html.parser")
            for res in soup.select(".result"):
                a = res.select_one(".result__title a") or res.select_one("a.result__a")
                if not a:
                    continue
                title = a.get_text(" ", strip=True)
                link = a.get("href") or ""
                if not title or not link:
                    continue
                norm = normalize_text(title)
                if norm in seen:
                    continue
                seen.add(norm)
                snippet_node = res.select_one(".result__snippet")
                snippet = snippet_node.get_text(" ", strip=True) if snippet_node else ""
                domain = urlparse(link).netloc.replace("www.", "") if link.startswith("http") else "No disponible"
                results.append({
                    "program": title,
                    "institution": domain or "No disponible",
                    "duration": "No disponible",
                    "modality": "No disponible",
                    "location": "No disponible",
                    "link": link,
                    "area": area,
                    "source_file": "Búsqueda web",
                    "type": "Externo",
                    "content": snippet,
                })
                if len(results) >= max_items:
                    return pd.DataFrame(results)
        except Exception:
            continue
    return pd.DataFrame(results)


def build_question_answer(question, analysis):
    strengths = rank_strengths(analysis["indicators"], top_n=3)
    weaknesses = rank_weaknesses(analysis["indicators"], top_n=3)
    best_strengths = ", ".join(strengths["indicator"].tolist()) or "fortalezas no disponibles"
    main_weaknesses = ", ".join(weaknesses["indicator"].tolist()) or "brechas no disponibles"
    profile = analysis["profile"]
    group = QUESTION_TO_GROUP.get(question, "General")
    if group == "Potencial y proyección":
        return f"Con el perfil actual (**{profile}**) y la combinación de fortalezas ({best_strengths}), el potencial de crecimiento parece razonable si se cierran primero las brechas más limitantes ({main_weaknesses})."
    if group == "Fortalezas y especialización":
        return f"Las áreas con mayor probabilidad de consolidarse como línea de especialización son aquellas donde el técnico ya muestra mejor desempeño relativo: {best_strengths}. Conviene reforzarlas con formación avanzada sin perder equilibrio global."
    if group == "Brechas y riesgos":
        return f"Las principales brechas del caso se concentran en {main_weaknesses}. Si estas áreas no se corrigen antes de forzar una especialización avanzada, existe riesgo de crear un perfil desequilibrado."
    return f"El plan de acción debería partir del nivel actual (**{profile}**) y priorizar la mejora de {main_weaknesses}. Después tendría sentido acelerar el desarrollo sobre fortalezas como {best_strengths}."


def build_executive_report(analysis, selected_species, selected_subspecies):
    global_summary = analysis["global"]
    trunks_df = analysis["trunks"].copy()
    strengths = rank_strengths(analysis["indicators"], top_n=3)
    weaknesses = rank_weaknesses(analysis["indicators"], top_n=3)
    priority_areas = build_priority_areas(analysis["indicators"], top_n=3)
    best_trunk = trunks_df.sort_values("vs_goal", ascending=False).iloc[0]["tronco"] if not trunks_df.empty else "No disponible"
    weak_trunk = trunks_df.sort_values("vs_goal", ascending=True).iloc[0]["tronco"] if not trunks_df.empty else "No disponible"
    strength_names = ", ".join(str(x) for x in strengths["indicator"].tolist()) or "No disponible"
    weakness_names = ", ".join(str(x) for x in weaknesses["indicator"].tolist()) or "No disponible"
    area_names = ", ".join(str(x) for x in priority_areas["tronco"].tolist()) or "No disponible"
    lines = [
        f"1. El técnico evaluado es **{analysis['name']}** y el perfil global propuesto es **{analysis['profile']}**.",
        f"2. El score medio global del assessment es **{global_summary['score_raw_avg']:.2f} / 6**.",
        f"3. El resultado se sitúa en **{format_pct(global_summary['vs_goal'])}** respecto al objetivo y en **{format_pct(global_summary['vs_bbdd'])}** respecto al benchmark.",
        f"4. El tronco más sólido es **{best_trunk}** y el tronco con mayor necesidad de desarrollo es **{weak_trunk}**.",
        f"5. Las fortalezas más visibles se concentran en: **{strength_names}**.",
        f"6. Las debilidades que más condicionan el desarrollo actual son: **{weakness_names}**.",
        f"7. Las áreas prioritarias del plan de carrera deberían centrarse primero en: **{area_names}**.",
        f"8. La lectura del perfil sugiere combinar formación base para cerrar brechas y formación especializada para consolidar ventajas competitivas.",
        f"9. El plan se ha elaborado para la especie **{selected_species}** y la subespecie **{selected_subspecies}**.",
        f"10. Se recomienda revisar de nuevo los indicadores prioritarios en la próxima evaluación para confirmar evolución y ajustar el itinerario.",
    ]
    return "\n".join(lines)


def render_course_lines(df, title):
    lines = [f"### {title}"]
    if df.empty:
        lines.append("- No hay cursos disponibles.")
        return lines
    for _, row in df.iterrows():
        url = row.get("link", "No disponible")
        url = url if isinstance(url, str) and url.startswith("http") else "No disponible"
        lines.append(
            f"- [{row.get('area', 'General')}] {row.get('program', 'Curso')} | {row.get('institution', 'No disponible')} | {row.get('duration', 'No disponible')} | {row.get('modality', 'No disponible')} | {row.get('location', 'No disponible')} | {url}"
        )
    return lines


def build_final_report_text(analysis, selected_species, selected_subspecies, base_internal, base_external, adv_internal, adv_external, career_plan_df, qa_items):
    global_summary = analysis["global"]
    trunks = analysis["trunks"].copy()
    strengths = rank_strengths(analysis["indicators"], top_n=5)
    weaknesses = rank_weaknesses(analysis["indicators"], top_n=5)
    priority_areas = build_priority_areas(analysis["indicators"], top_n=4)
    specialization_areas = build_specialization_areas(analysis["indicators"], top_n=3)

    lines = []
    lines.append("# Informe integrado de assessment y plan de carrera")
    lines.append("")
    lines.append("## 1. Resumen ejecutivo")
    lines.append(build_executive_report(analysis, selected_species, selected_subspecies))
    lines.append("")
    lines.append("## 2. Perfil del técnico")
    lines.append(f"- Nombre: {analysis['name']}")
    lines.append(f"- Especie seleccionada: {selected_species}")
    lines.append(f"- Subespecie seleccionada: {selected_subspecies}")
    lines.append(f"- Categoría propuesta: {analysis['profile']}")
    lines.append(f"- Score global medio: {global_summary['score_raw_avg']:.2f} / 6")
    lines.append(f"- Vs objetivo global: {format_pct(global_summary['vs_goal'])}")
    lines.append(f"- Vs benchmark global: {format_pct(global_summary['vs_bbdd'])}")
    if analysis.get("profile_ratio") is not None:
        lines.append(f"- Ratio de clasificación Excel: {analysis['profile_ratio']:.3f}")
        lines.append(f"- Ranking Excel: {analysis.get('profile_rank', 'No disponible')}")
    lines.append("")
    lines.append("## 3. Resultados por troncos")
    for _, row in trunks.iterrows():
        lines.append(f"- {canonical_trunk_name(row['tronco'])}: score medio {row['score_raw_avg']:.2f} | vs objetivo {format_pct(row['vs_goal'])} | vs benchmark {format_pct(row['vs_bbdd'])}")
    lines.append("")
    lines.append("## 4. Indicadores clave")
    lines.append("### Fortalezas")
    for _, row in strengths.iterrows():
        lines.append(f"- {row['indicator']} ({canonical_trunk_name(row['tronco'])}): score {row['score_raw']:.1f} | vs objetivo {format_pct(row['vs_goal'])} | vs benchmark {format_pct(row['vs_bbdd'])}")
    lines.append("### Debilidades")
    for _, row in weaknesses.iterrows():
        lines.append(f"- {row['indicator']} ({canonical_trunk_name(row['tronco'])}): score {row['score_raw']:.1f} | vs objetivo {format_pct(row['vs_goal'])} | vs benchmark {format_pct(row['vs_bbdd'])}")
    lines.append("")
    lines.append("## 5. Áreas prioritarias de formación")
    for _, row in priority_areas.iterrows():
        lines.append(f"- {canonical_trunk_name(row['tronco'])}: prioridad relativa {row['priority']:.2f}")
    lines.append("")
    lines.append("## 6. Áreas de especialización recomendadas")
    for _, row in specialization_areas.iterrows():
        lines.append(f"- {canonical_trunk_name(row['tronco'])}: potencial relativo {row['priority']:.2f}")
    lines.append("")
    lines.extend(render_career_plan_lines(career_plan_df))
    lines.append("")
    lines.append("## 8. Formación recomendada")
    lines.extend(render_course_lines(base_internal, "8.1 Formación base interna"))
    lines.extend(render_course_lines(base_external, "8.2 Formación base externa"))
    lines.extend(render_course_lines(adv_internal, "8.3 Formación especializada interna"))
    lines.extend(render_course_lines(adv_external, "8.4 Formación especializada externa"))
    lines.append("")
    lines.append("## 9. Gráficos de araña del assessment")
    lines.append("- Se incorporan al informe exportado un radar global por áreas y cuatro radares específicos: Alimentación, Sanidad, Manejo y Herramientas.")

    included_qas = [x for x in qa_items if x.get("include", True)]
    if included_qas:
        lines.append("")
        lines.append("## 10. Preguntas y respuestas incorporadas")
        for item in included_qas:
            lines.append(f"### {item['group']} — {item['question']}")
            lines.append(item['answer'])
            lines.append("")
    return "\n".join(lines)


def make_excel_export(analysis, base_internal, base_external, adv_internal, adv_external, career_plan_df, qa_items):
    bio = io.BytesIO()
    summary_df = pd.DataFrame([{
        "Nombre": analysis["name"],
        "Categoría": analysis["profile"],
        "Score global medio": round(analysis["global"]["score_raw_avg"], 2),
        "Vs objetivo global": pct(analysis["global"]["vs_goal"]),
        "Vs benchmark global": pct(analysis["global"]["vs_bbdd"]),
        "Ratio Excel": analysis.get("profile_ratio"),
        "Ranking Excel": analysis.get("profile_rank"),
    }])

    trunks_df = analysis["trunks"].copy()
    trunks_df["tronco"] = trunks_df["tronco"].apply(canonical_trunk_name)
    trunks_df["vs_goal"] = trunks_df["vs_goal"].apply(pct)
    trunks_df["vs_bbdd"] = trunks_df["vs_bbdd"].apply(pct)
    trunks_df["vs_max"] = trunks_df["vs_max"].apply(pct)

    indicators_df = analysis["indicators"].copy()
    indicators_df["tronco"] = indicators_df["tronco"].apply(canonical_trunk_name)
    indicators_df["vs_goal"] = indicators_df["vs_goal"].apply(pct)
    indicators_df["vs_bbdd"] = indicators_df["vs_bbdd"].apply(pct)
    indicators_df["vs_max"] = indicators_df["vs_max"].apply(pct)

    qa_df = pd.DataFrame(qa_items) if qa_items else pd.DataFrame(columns=["group", "question", "answer", "include"])

    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        summary_df.to_excel(writer, sheet_name="Resumen", index=False)
        trunks_df.to_excel(writer, sheet_name="Troncos", index=False)
        indicators_df.to_excel(writer, sheet_name="Indicadores", index=False)
        career_plan_df.to_excel(writer, sheet_name="Plan carrera", index=False)
        base_internal.to_excel(writer, sheet_name="Base interna", index=False)
        base_external.to_excel(writer, sheet_name="Base externa", index=False)
        adv_internal.to_excel(writer, sheet_name="Especializacion int", index=False)
        adv_external.to_excel(writer, sheet_name="Especializacion ext", index=False)
        qa_df.to_excel(writer, sheet_name="Preguntas", index=False)
    bio.seek(0)
    return bio.getvalue()


def make_docx_bytes(report_text, chart_images=None):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    inserted_charts = False
    for line in report_text.splitlines():
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            if heading == "Gráficos de araña del assessment" and chart_images:
                for title, image_bytes in chart_images:
                    doc.add_paragraph(title)
                    if image_bytes:
                        bio_img = io.BytesIO(image_bytes)
                        doc.add_picture(bio_img, width=Inches(5.8))
                    else:
                        doc.add_paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)")
                inserted_charts = True
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    if chart_images and not inserted_charts:
        doc.add_heading("Gráficos de araña del assessment", level=1)
        for title, image_bytes in chart_images:
            doc.add_paragraph(title)
            if image_bytes:
                bio_img = io.BytesIO(image_bytes)
                doc.add_picture(bio_img, width=Inches(5.8))
            else:
                doc.add_paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def make_pdf_bytes(report_text, chart_images=None):
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#1f2937"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#1f2937"))
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#111827"))
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, bulletIndent=0)

    story = []
    inserted_charts = False
    for line in report_text.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        elif stripped.startswith("# "):
            story.append(Paragraph(escape(stripped[2:].strip()), h1))
            story.append(Spacer(1, 8))
        elif stripped.startswith("## "):
            heading = stripped[3:].strip()
            story.append(Paragraph(escape(heading), h2))
            story.append(Spacer(1, 6))
            if heading == "Gráficos de araña del assessment" and chart_images:
                for title, image_bytes in chart_images:
                    story.append(Paragraph(escape(title), h3))
                    if image_bytes:
                        img = RLImage(io.BytesIO(image_bytes), width=500, height=380)
                        story.append(img)
                    else:
                        story.append(Paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)", body))
                    story.append(Spacer(1, 8))
                inserted_charts = True
        elif stripped.startswith("### "):
            story.append(Paragraph(escape(stripped[4:].strip()), h3))
            story.append(Spacer(1, 4))
        elif stripped.startswith("- "):
            story.append(Paragraph("• " + escape(stripped[2:].strip()), bullet))
        else:
            story.append(Paragraph(escape(stripped), body))

    if chart_images and not inserted_charts:
        story.append(Paragraph("Gráficos de araña del assessment", h2))
        story.append(Spacer(1, 6))
        for title, image_bytes in chart_images:
            story.append(Paragraph(escape(title), h3))
            if image_bytes:
                img = RLImage(io.BytesIO(image_bytes), width=500, height=380)
                story.append(img)
            else:
                story.append(Paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)", body))
            story.append(Spacer(1, 8))

    doc.build(story)
    bio.seek(0)
    return bio.getvalue()


def clean_url(url):
    url = str(url).strip() if url not in (None, "") else ""
    return url if url.startswith("http://") or url.startswith("https://") else None


def display_course_table(title, df):
    st.markdown(f"### {title}")
    if df.empty:
        st.info("No hay cursos disponibles para este bloque.")
        return
    view = df.copy()
    view["Abrir enlace"] = view["link"].apply(clean_url)
    keep_cols = [col for col in ["area", "program", "institution", "type", "duration", "modality", "location", "source_file", "Abrir enlace"] if col in view.columns]
    view = view[keep_cols].rename(columns={
        "area": "Área",
        "program": "Programa",
        "institution": "Institución",
        "type": "Tipo",
        "duration": "Duración",
        "modality": "Modalidad",
        "location": "Ubicación",
        "source_file": "Fuente",
    })
    st.dataframe(
        view,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Abrir enlace": st.column_config.LinkColumn("Enlace", display_text="Abrir"),
        },
    )


def build_catalog_summary(catalog_df):
    if catalog_df.empty:
        return pd.DataFrame(columns=["Catálogo", "Cursos detectados", "Especies detectadas"])
    rows = []
    for source, grp in catalog_df.groupby("source_file"):
        species = sorted(set(str(x) for x in grp["species"].dropna().tolist() if str(x).strip()))
        rows.append({
            "Catálogo": source,
            "Cursos detectados": int(len(grp)),
            "Especies detectadas": ", ".join(species[:6]) if species else "General",
        })
    return pd.DataFrame(rows).sort_values(["Cursos detectados", "Catálogo"], ascending=[False, True])


def pretty_profile_label(label):
    n = normalize_text(label)
    mapping = {
        "basico": "Básico",
        "controla": "Controla",
        "supera": "Supera",
        "certificado": "Certificado",
        "excelente": "Excelente",
        "master": "Máster",
        "maximo": "Máximo",
    }
    return mapping.get(n, str(label))


def normalize_species_for_matrix(species):
    n = normalize_text(species)
    if "cunic" in n or "conej" in n:
        return "Cunicultura"
    if "avic" in n:
        return "Avicultura"
    if "porcin" in n:
        return "Porcino"
    if "vacuno" in n and ("lech" in n or "leche" in n):
        return "Vacuno de leche"
    if "vacuno" in n and "carn" in n:
        return "Vacuno de carne"
    if "ovino" in n or "caprino" in n or "peque" in n:
        return "Ovino y caprino"
    return species


def normalize_trunk_for_matrix(trunk):
    canon = canonical_trunk_name(trunk)
    mapping = {
        "Alimentación": "Nutrición",
        "Sanidad": "Sanidad",
        "Manejo": "Manejo",
        "Herramientas": "Herramientas",
    }
    return mapping.get(canon, canon)


@st.cache_data(show_spinner=False)
def load_module_matrix(path_str):
    return pd.read_excel(path_str, sheet_name="Matriz módulos").fillna("")


@st.cache_data(show_spinner=False)
def load_assignment_rules(path_str):
    try:
        return pd.read_excel(path_str, sheet_name="Reglas de asignación").fillna("")
    except Exception:
        return pd.DataFrame()


def build_module_matrix_df(asset_paths):
    matrix_path = asset_paths.get("module_matrix")
    if matrix_path and Path(matrix_path).exists():
        return load_module_matrix(str(matrix_path))
    return pd.DataFrame()


def build_assignment_rules_df(asset_paths):
    matrix_path = asset_paths.get("module_matrix")
    if matrix_path and Path(matrix_path).exists():
        return load_assignment_rules(str(matrix_path))
    return pd.DataFrame()


def trunk_profile_from_row(trunk_row, thresholds):
    info = classify_profile_from_ratio(trunk_row.get("vs_max"), thresholds)
    label = pretty_profile_label(info.get("label", "No disponible"))
    return label, info.get("ratio")


def recommended_module_level(trunk_profile, strategic=False):
    n = normalize_text(trunk_profile)
    if n in {"basico", "controla", "supera"}:
        return "Certificado", "Prioritario", "El resultado del tronco está por debajo de Certificado; conviene consolidar base operativa."
    if n == "certificado":
        if strategic:
            return "Máster", "Recomendado", "El tronco ya está en Certificado y se considera estratégico para el rol; conviene avanzar a Máster."
        return "Certificado", "Mantenimiento", "El tronco ya está en Certificado; conviene consolidar aplicación y evidencia antes de escalar."
    if n == "excelente":
        return "Máster", "Recomendado", "El tronco muestra alto potencial; conviene especialización de nivel Máster."
    if n == "master":
        return "Máster", "Mentoring", "El técnico ya está en Máster; debe orientarse a mentoring, docencia interna y casos complejos."
    return "Certificado", "Recomendado", "Se propone Certificado como itinerario por defecto mientras se valida el caso."


def build_career_plan_df(analysis, selected_species, module_matrix_df, strategic_trunks=None):
    strategic_trunks = strategic_trunks or []
    if module_matrix_df.empty:
        return pd.DataFrame()

    species_matrix = normalize_species_for_matrix(selected_species)
    thresholds = analysis.get("profile_thresholds", [])
    rows = []
    for _, trunk_row in analysis["trunks"].iterrows():
        trunk_display = canonical_trunk_name(trunk_row["tronco"])
        matrix_trunk = normalize_trunk_for_matrix(trunk_row["tronco"])
        trunk_profile, trunk_ratio = trunk_profile_from_row(trunk_row, thresholds)
        strategic = matrix_trunk in strategic_trunks
        level, priority, rationale = recommended_module_level(trunk_profile, strategic=strategic)

        subset = module_matrix_df[
            (module_matrix_df["Especie"].apply(lambda x: normalize_text(x) == normalize_text(species_matrix)))
            & (module_matrix_df["Tronco assessment"].apply(lambda x: normalize_text(x) == normalize_text(matrix_trunk)))
            & (module_matrix_df["Nivel"].apply(lambda x: normalize_text(x) == normalize_text(level)))
        ]
        module = subset.iloc[0] if not subset.empty else None

        rows.append({
            "Tronco": trunk_display,
            "Resultado tronco": trunk_profile,
            "Ratio vs max": round(trunk_ratio, 3) if trunk_ratio is not None else None,
            "Estratégico": "Sí" if strategic else "No",
            "Nivel recomendado": level,
            "Prioridad": priority,
            "Código módulo": module["Código módulo"] if module is not None else "No disponible",
            "Objetivo": module["Objetivo"] if module is not None else "No disponible",
            "Perfil destinatario": module["Perfil destinatario"] if module is not None else "No disponible",
            "Contenidos clave": module["Contenidos clave"] if module is not None else "No disponible",
            "Formato recomendado": module["Formato recomendado"] if module is not None else "No disponible",
            "Evaluación / evidencia": module["Evaluación / evidencia"] if module is not None else "No disponible",
            "Carga estimada": module["Carga estimada"] if module is not None else "No disponible",
            "Fuente curricular base": module["Fuente curricular base"] if module is not None else "No disponible",
            "Criterio de acceso sugerido": module["Criterio de acceso sugerido"] if module is not None else rationale,
            "Justificación directiva": rationale,
        })
    return pd.DataFrame(rows)


def career_plan_summary(career_plan_df):
    if career_plan_df.empty:
        return {"certificado": 0, "master": 0, "mentoring": 0}
    return {
        "certificado": int((career_plan_df["Nivel recomendado"] == "Certificado").sum()),
        "master": int((career_plan_df["Nivel recomendado"] == "Máster").sum()),
        "mentoring": int((career_plan_df["Prioridad"] == "Mentoring").sum()),
    }


def render_career_plan_lines(career_plan_df):
    lines = ["## 7. Recomendación de módulos del plan de carrera"]
    if career_plan_df.empty:
        lines.append("- No se ha podido construir la matriz de módulos recomendados.")
        return lines
    for _, row in career_plan_df.iterrows():
        lines.append(
            f"- {row['Tronco']}: resultado {row['Resultado tronco']} | nivel recomendado {row['Nivel recomendado']} | prioridad {row['Prioridad']} | módulo {row['Código módulo']}"
        )
        lines.append(f"  Objetivo: {row['Objetivo']}")
        lines.append(f"  Contenidos: {row['Contenidos clave']}")
        lines.append(f"  Justificación: {row['Justificación directiva']}")
    return lines


def display_career_plan_table(career_plan_df):
    if career_plan_df.empty:
        st.info("No se ha podido construir la recomendación de módulos del plan de carrera.")
        return
    view = career_plan_df[[
        "Tronco", "Resultado tronco", "Estratégico", "Nivel recomendado", "Prioridad", "Código módulo", "Carga estimada"
    ]].copy()
    st.dataframe(view, use_container_width=True, hide_index=True)
    for _, row in career_plan_df.iterrows():
        with st.expander(f"{row['Tronco']} · {row['Nivel recomendado']} · {row['Código módulo']}"):
            st.markdown(f"**Objetivo**: {row['Objetivo']}")
            st.markdown(f"**Perfil destinatario**: {row['Perfil destinatario']}")
            st.markdown(f"**Contenidos clave**: {row['Contenidos clave']}")
            st.markdown(f"**Formato recomendado**: {row['Formato recomendado']}")
            st.markdown(f"**Evaluación / evidencia**: {row['Evaluación / evidencia']}")
            st.markdown(f"**Carga estimada**: {row['Carga estimada']}")
            st.markdown(f"**Fuente curricular base**: {row['Fuente curricular base']}")
            st.markdown(f"**Justificación directiva**: {row['Justificación directiva']}")


def reset_app():
    st.session_state["reset_counter"] = st.session_state.get("reset_counter", 0) + 1
    st.session_state["qa_items"] = []
    st.session_state["final_report_text"] = ""
    st.rerun()



import base64
from datetime import datetime
import streamlit.components.v1 as components

APP_PASSWORD = "TechTeam2026+"
ASSET_CANDIDATES = {
    "nutreco": ["Logo Nutreco.jpg"],
    "techteam": ["Logo TechTeam 2.jpg"],
    "strip": ["Solapa rosa.jpg"],
    "finder": ["buscador_formacion_produccion_animal_v2.html", "buscador_formacion_produccion_animal.html"],
    "module_matrix": ["Matriz_modulos_plan_formacion_tecnica.xlsx"],
}


def get_asset_path(candidates):
    search_roots = [Path('.'), Path('/mnt/data')]
    for root in search_roots:
        for name in candidates:
            p = root / name
            if p.exists():
                return p
    return None


def get_asset_paths():
    return {key: get_asset_path(names) for key, names in ASSET_CANDIDATES.items()}


def image_to_data_uri(path):
    if not path or not Path(path).exists():
        return ""
    data = Path(path).read_bytes()
    ext = Path(path).suffix.lower()
    mime = "image/png" if ext == ".png" else "image/jpeg"
    return f"data:{mime};base64," + base64.b64encode(data).decode("ascii")


def display_corporate_header(asset_paths):
    strip_path = asset_paths.get("strip")
    if strip_path and Path(strip_path).exists():
        st.image(str(strip_path), use_container_width=True)
    c1, c2, c3 = st.columns([1.0, 1.8, 1.2])
    if asset_paths.get("nutreco"):
        with c1:
            st.image(str(asset_paths["nutreco"]), use_container_width=True)
    with c2:
        st.markdown(
            """
            <div style="padding-top:8px;">
              <div style="font-size:34px;font-weight:800;color:#143b8f;line-height:1.05;">Assessment + Plan de Carrera Técnico</div>
              <div style="margin-top:8px;color:#6b7280;font-size:15px;line-height:1.45;">
                Entorno corporativo TechTeam · Nutreco.
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    if asset_paths.get("techteam"):
        with c3:
            st.image(str(asset_paths["techteam"]), use_container_width=True)


def require_password(asset_paths):
    if st.session_state.get("auth_ok"):
        return True
    display_corporate_header(asset_paths)
    st.markdown("### Acceso restringido")
    st.write("Introduce la contraseña para acceder a la aplicación.")
    entered = st.text_input("Contraseña", type="password", key="auth_password")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("Acceder", type="primary", use_container_width=True):
            if entered == APP_PASSWORD:
                st.session_state["auth_ok"] = True
                st.session_state.pop("auth_error", None)
                st.rerun()
            else:
                st.session_state["auth_error"] = "Contraseña incorrecta."
    with col2:
        st.caption("La contraseña se solicita al inicio de cada sesión del navegador.")
    if st.session_state.get("auth_error"):
        st.error(st.session_state["auth_error"])
    st.stop()


def load_finder_html(asset_paths):
    finder_path = asset_paths.get("finder")
    if finder_path and Path(finder_path).exists():
        return Path(finder_path).read_text(encoding="utf-8")
    return None


def markdownish_to_html(report_text):
    html_parts = []
    in_list = False
    for raw_line in report_text.splitlines():
        line = raw_line.strip()
        if not line:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            continue
        if line.startswith("# "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h1>{escape(line[2:].strip())}</h1>")
        elif line.startswith("## "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h2>{escape(line[3:].strip())}</h2>")
        elif line.startswith("### "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h3>{escape(line[4:].strip())}</h3>")
        elif line.startswith("- "):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{escape(line[2:].strip())}</li>")
        else:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<p>{escape(line)}</p>")
    if in_list:
        html_parts.append("</ul>")
    return "\n".join(html_parts)


def charts_to_html(chart_images):
    if not chart_images:
        return ""
    blocks = []
    for title, image_bytes in chart_images:
        if not image_bytes:
            continue
        b64 = base64.b64encode(image_bytes).decode("ascii")
        blocks.append(
            f"<div class='chart-block'><h3>{escape(title)}</h3><img src='data:image/png;base64,{b64}' alt='{escape(title)}'></div>"
        )
    if not blocks:
        return ""
    return "<div class='chart-grid'>" + "".join(blocks) + "</div>"


def make_html_report_bytes(report_text, analysis, selected_species, selected_subspecies, chart_images=None, asset_paths=None):
    asset_paths = asset_paths or get_asset_paths()
    logos = {k: image_to_data_uri(v) for k, v in asset_paths.items() if k in {"nutreco", "techteam", "strip"}}
    body_html = markdownish_to_html(report_text)
    chart_html = charts_to_html(chart_images)
    profile = analysis.get("profile", "No disponible")
    global_score = analysis.get("global", {}).get("score_raw_avg")
    global_score_txt = "No disponible" if global_score is None or pd.isna(global_score) else f"{global_score:.2f} / 6"

    html_doc = f"""<!DOCTYPE html>
<html lang='es'>
<head>
<meta charset='utf-8'>
<meta name='viewport' content='width=device-width, initial-scale=1'>
<title>Informe plan de carrera</title>
<style>
:root {{
  --nutreco-blue:#143b8f;
  --nanta-red:#ef233c;
  --pink:#d81b90;
  --line:#dbe3ef;
  --muted:#6b7280;
  --bg:#f7f9fc;
}}
* {{ box-sizing:border-box; }}
body {{ margin:0; font-family:Arial, Helvetica, sans-serif; color:#1f2937; background:var(--bg); }}
.top-strip {{ height:28px; background-image:url('{logos.get("strip", "")}'); background-size:cover; background-position:center; }}
.container {{ max-width:1200px; margin:0 auto; padding:22px; }}
.hero {{ background:linear-gradient(135deg, rgba(20,59,143,.98), rgba(216,27,144,.93)); color:#fff; border-radius:24px; padding:20px 22px; }}
.hero-grid {{ display:grid; grid-template-columns:220px 260px 1fr; gap:14px; align-items:center; }}
.hero-grid img {{ max-width:100%; max-height:96px; object-fit:contain; background:#fff; border-radius:14px; padding:8px; }}
.hero-title h1 {{ margin:0 0 8px; font-size:30px; line-height:1.05; }}
.hero-title p {{ margin:0; font-size:14px; opacity:.97; }}
.meta {{ display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:12px; margin-top:18px; }}
.meta .box {{ background:#fff; border:1px solid var(--line); border-radius:16px; padding:12px; }}
.meta small {{ display:block; color:var(--muted); margin-bottom:4px; }}
.meta strong {{ color:var(--nutreco-blue); font-size:18px; }}
.section {{ background:#fff; border:1px solid var(--line); border-radius:20px; padding:18px; margin-top:18px; }}
.section h1 {{ font-size:26px; color:var(--nutreco-blue); margin:0 0 12px; }}
.section h2 {{ font-size:20px; color:var(--nutreco-blue); border-left:8px solid var(--nanta-red); padding-left:12px; margin:22px 0 10px; }}
.section h3 {{ font-size:16px; color:#111827; margin:16px 0 8px; }}
.section p, .section li {{ font-size:14px; line-height:1.55; }}
.chart-grid {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:16px; margin-top:14px; }}
.chart-block {{ border:1px solid var(--line); border-radius:16px; padding:12px; background:#fff; }}
.chart-block img {{ width:100%; height:auto; display:block; }}
.footer-note {{ margin-top:18px; color:var(--muted); font-size:12px; }}
@media (max-width:1000px) {{ .hero-grid, .meta, .chart-grid {{ grid-template-columns:1fr; }} }}
</style>
</head>
<body>
<div class='top-strip'></div>
<div class='container'>
  <div class='hero'>
    <div class='hero-grid'>
      <img src='{logos.get("nutreco", "")}' alt='Nutreco'>
      <img src='{logos.get("techteam", "")}' alt='TechTeam'>
      <div class='hero-title'>
        <h1>Informe corporativo · Assessment + Plan de Carrera Técnico</h1>
        <p>Documento generado para revisión técnica y planificación del desarrollo profesional.</p>
      </div>
    </div>
  </div>
  <div class='meta'>
    <div class='box'><small>Técnico</small><strong>{escape(str(analysis.get("name", "No disponible")))}</strong></div>
    <div class='box'><small>Perfil</small><strong>{escape(profile)}</strong></div>
    <div class='box'><small>Especie / subespecie</small><strong>{escape(selected_species)} · {escape(selected_subspecies)}</strong></div>
    <div class='box'><small>Score global</small><strong>{escape(global_score_txt)}</strong></div>
  </div>
  <div class='section'>
    {body_html}
    {chart_html}
    <div class='footer-note'>Generado el {escape(datetime.now().strftime('%d/%m/%Y %H:%M'))}. Documento interno con plantilla corporativa.</div>
  </div>
</div>
</body>
</html>"""
    return html_doc.encode("utf-8")


def make_docx_bytes(report_text, chart_images=None, asset_paths=None):
    asset_paths = asset_paths or get_asset_paths()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title_p = doc.add_paragraph()
    for key in ["catedra", "nutreco", "techteam"]:
        p = asset_paths.get(key)
        if p and Path(p).exists():
            try:
                title_p.add_run().add_picture(str(p), width=Inches(1.3 if key != 'techteam' else 1.8))
                title_p.add_run("   ")
            except Exception:
                pass
    doc.add_heading("Assessment + Plan de Carrera Técnico", level=0)

    inserted_charts = False
    for line in report_text.splitlines():
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            if heading == "Gráficos de araña del assessment" and chart_images:
                for title, image_bytes in chart_images:
                    doc.add_paragraph(title)
                    if image_bytes:
                        bio_img = io.BytesIO(image_bytes)
                        doc.add_picture(bio_img, width=Inches(5.8))
                    else:
                        doc.add_paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)")
                inserted_charts = True
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    if chart_images and not inserted_charts:
        doc.add_heading("Gráficos de araña del assessment", level=1)
        for title, image_bytes in chart_images:
            doc.add_paragraph(title)
            if image_bytes:
                bio_img = io.BytesIO(image_bytes)
                doc.add_picture(bio_img, width=Inches(5.8))
            else:
                doc.add_paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def make_pdf_bytes(report_text, chart_images=None, asset_paths=None):
    asset_paths = asset_paths or get_asset_paths()
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#143b8f"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#143b8f"))
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#111827"))
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, bulletIndent=0)

    story = []
    for key, width in [("catedra", 130), ("nutreco", 110), ("techteam", 160)]:
        p = asset_paths.get(key)
        if p and Path(p).exists():
            try:
                story.append(RLImage(str(p), width=width, height=width * 0.45))
                story.append(Spacer(1, 6))
            except Exception:
                pass
    story.append(Paragraph("Assessment + Plan de Carrera Técnico", h1))
    story.append(Spacer(1, 8))

    inserted_charts = False
    for line in report_text.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        elif stripped.startswith("# "):
            story.append(Paragraph(escape(stripped[2:].strip()), h1))
            story.append(Spacer(1, 8))
        elif stripped.startswith("## "):
            heading = stripped[3:].strip()
            story.append(Paragraph(escape(heading), h2))
            story.append(Spacer(1, 6))
            if heading == "Gráficos de araña del assessment" and chart_images:
                for title, image_bytes in chart_images:
                    story.append(Paragraph(escape(title), h3))
                    if image_bytes:
                        story.append(RLImage(io.BytesIO(image_bytes), width=500, height=380))
                    else:
                        story.append(Paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)", body))
                    story.append(Spacer(1, 8))
                inserted_charts = True
        elif stripped.startswith("### "):
            story.append(Paragraph(escape(stripped[4:].strip()), h3))
            story.append(Spacer(1, 4))
        elif stripped.startswith("- "):
            story.append(Paragraph("• " + escape(stripped[2:].strip()), bullet))
        else:
            story.append(Paragraph(escape(stripped), body))

    if chart_images and not inserted_charts:
        story.append(Paragraph("Gráficos de araña del assessment", h2))
        story.append(Spacer(1, 6))
        for title, image_bytes in chart_images:
            story.append(Paragraph(escape(title), h3))
            if image_bytes:
                story.append(RLImage(io.BytesIO(image_bytes), width=500, height=380))
            else:
                story.append(Paragraph("(No se pudo generar la imagen del gráfico en esta ejecución.)", body))
            story.append(Spacer(1, 8))

    doc.build(story)
    bio.seek(0)
    return bio.getvalue()


asset_paths = get_asset_paths()
require_password(asset_paths)

if "qa_items" not in st.session_state:
    st.session_state["qa_items"] = []

st.markdown(
    "<style>div[data-testid='stTabs'] button {font-weight:700;} .stDownloadButton button {font-weight:700;}</style>",
    unsafe_allow_html=True,
)

display_corporate_header(asset_paths)
st.caption(
    "Carga la evaluación de un técnico, interpreta el assessment, genera un informe integrado y añade una pestaña específica de plan de carrera con recomendación de módulos por especie, tronco y nivel."
)

sidebar = st.sidebar
sidebar.header("Carga y configuración")
reset_key = st.session_state.get("reset_counter", 0)
assessment_file = sidebar.file_uploader("Assessment del técnico (.xlsm / .xlsx)", type=["xlsm", "xlsx"], key=f"assessment_{reset_key}")
manual_catalogs = sidebar.file_uploader("Catálogos de formación (.xlsx) — opcional, puedes subir varios", type=["xlsx", "xlsm"], accept_multiple_files=True, key=f"catalogs_{reset_key}")
internal_courses_per_area = sidebar.number_input("Cursos internos por área", min_value=1, max_value=10, value=3, step=1)
external_courses_per_area = sidebar.number_input("Cursos externos por área", min_value=1, max_value=10, value=2, step=1)
use_web = sidebar.checkbox("Refuerza con búsqueda web externa", value=True)

with sidebar.expander("Clasificación del perfil", expanded=False):
    st.caption(
        "La categoría global se toma del propio Excel y debe coincidir con ENLACES DATOS!N44. La lógica usa ENLACES DATOS!K37 como ratio y los cortes de ENLACES DATOS!M56:M61."
    )

if sidebar.button("Nuevo / borrar evaluación cargada", use_container_width=True):
    reset_app()

assessment_tab, career_tab, finder_tab = st.tabs(["Assessment", "Plan de carrera", "Buscador global de formación"])

analysis = None
catalog_df = pd.DataFrame()
selected_species = None
selected_subspecies = None
base_internal = pd.DataFrame()
base_external = pd.DataFrame()
adv_internal = pd.DataFrame()
adv_external = pd.DataFrame()
priority_areas = pd.DataFrame()
specialization_areas = pd.DataFrame()
report_charts = []
module_matrix_df = build_module_matrix_df(asset_paths)
rules_df = build_assignment_rules_df(asset_paths)
career_plan_df = pd.DataFrame()

if assessment_file:
    try:
        analysis = parse_assessment(assessment_file)
    except Exception as exc:
        st.error(f"No se ha podido procesar el assessment: {exc}")
        st.stop()

    catalog_frames = []
    manual_names = []
    for file in manual_catalogs or []:
        try:
            catalog_frames.append(parse_uploaded_catalog(file))
            manual_names.append(file.name)
        except Exception as exc:
            st.warning(f"No se ha podido leer el catálogo {file.name}: {exc}")

    for path in autodiscover_catalogs(exclude_names=manual_names):
        try:
            catalog_frames.append(load_catalog_from_path(str(path)))
        except Exception:
            continue

    catalog_df = pd.concat([df for df in catalog_frames if not df.empty], ignore_index=True) if catalog_frames else pd.DataFrame()

    species_options = sorted(set([analysis["species_from_file"]] + ([] if catalog_df.empty else [x for x in catalog_df["species"].dropna().unique().tolist() if str(x).strip()])))
    default_species = analysis["species_from_file"]
    selected_species = assessment_tab.selectbox("Especie", species_options if species_options else [default_species], index=0, key="species_select")

    subspecies_options = ["General"]
    if not catalog_df.empty and "subspecies" in catalog_df.columns:
        subs = catalog_df[catalog_df["species"].apply(lambda x: species_match(selected_species, x))]["subspecies"].dropna().astype(str).tolist()
        subspecies_options = sorted(set(["General"] + [x for x in subs if x.strip()]))
    selected_subspecies = assessment_tab.selectbox("Subespecie", subspecies_options, index=0, key="subspecies_select")

    indicators_df = analysis["indicators"].copy()
    priority_areas = build_priority_areas(indicators_df, top_n=4)
    specialization_areas = build_specialization_areas(indicators_df, top_n=3)

    base_internal_list, adv_internal_list, base_external_list, adv_external_list = [], [], [], []
    for _, row in priority_areas.iterrows():
        rec = recommend_internal_courses(catalog_df, selected_species, selected_subspecies, row["tronco"], max_items=internal_courses_per_area)
        if not rec.empty:
            base_internal_list.append(rec)
        if use_web:
            ext = web_search_courses(selected_species, row["tronco"], selected_subspecies, max_items=external_courses_per_area)
            if not ext.empty:
                base_external_list.append(ext)
    for _, row in specialization_areas.iterrows():
        rec = recommend_internal_courses(catalog_df, selected_species, selected_subspecies, row["tronco"], max_items=internal_courses_per_area)
        if not rec.empty:
            adv_internal_list.append(rec)
        if use_web:
            ext = web_search_courses(selected_species, row["tronco"], selected_subspecies, max_items=external_courses_per_area)
            if not ext.empty:
                adv_external_list.append(ext)

    base_internal = pd.concat(base_internal_list, ignore_index=True).drop_duplicates(subset=["program", "institution", "source_file"]) if base_internal_list else pd.DataFrame(columns=["program", "institution", "duration", "modality", "location", "link", "area", "source_file", "type"])
    adv_internal = pd.concat(adv_internal_list, ignore_index=True).drop_duplicates(subset=["program", "institution", "source_file"]) if adv_internal_list else pd.DataFrame(columns=["program", "institution", "duration", "modality", "location", "link", "area", "source_file", "type"])
    base_external = pd.concat(base_external_list, ignore_index=True).drop_duplicates(subset=["program", "institution", "source_file"]) if base_external_list else pd.DataFrame(columns=["program", "institution", "duration", "modality", "location", "link", "area", "source_file", "type"])
    adv_external = pd.concat(adv_external_list, ignore_index=True).drop_duplicates(subset=["program", "institution", "source_file"]) if adv_external_list else pd.DataFrame(columns=["program", "institution", "duration", "modality", "location", "link", "area", "source_file", "type"])

    default_strategic = [normalize_trunk_for_matrix(x) for x in specialization_areas["tronco"].tolist()]
    default_strategic = [x for x in [*dict.fromkeys(default_strategic)] if x in ["Nutrición", "Sanidad", "Manejo", "Herramientas"]]
    with career_tab:
        st.subheader("Plan de carrera por módulos")
        if module_matrix_df.empty:
            st.warning("No se ha localizado la matriz de módulos del plan de carrera. Añade el archivo 'Matriz_modulos_plan_formacion_tecnica.xlsx'.")
        else:
            strategic_trunks = st.multiselect(
                "Troncos estratégicos para el rol actual o futuro del técnico",
                options=["Nutrición", "Sanidad", "Manejo", "Herramientas"],
                default=default_strategic,
                help="Si un tronco está en Certificado y además es estratégico para su función, la recomendación escalará a Máster, según la regla directiva.",
            )
            if not rules_df.empty:
                st.markdown("### Regla directiva de asignación")
                st.dataframe(rules_df, use_container_width=True, hide_index=True)
            career_plan_df = build_career_plan_df(analysis, selected_species, module_matrix_df, strategic_trunks=strategic_trunks)
            summary = career_plan_summary(career_plan_df)
            c1, c2, c3 = st.columns(3)
            c1.metric("Módulos Certificado", summary["certificado"])
            c2.metric("Módulos Máster", summary["master"])
            c3.metric("Troncos en mentoring", summary["mentoring"])
            st.caption("La lógica de recomendación sigue la arquitectura Certificado / Máster por especie y tronco, usando la matriz corporativa cargada en la app.")
            display_career_plan_table(career_plan_df)

    report_charts = build_report_chart_pack(analysis)

with assessment_tab:
    if analysis is None:
        st.info("Sube un archivo de assessment para activar el análisis, el plan de carrera y las exportaciones.")
    else:
        if not catalog_df.empty:
            st.subheader("Resumen de catálogos cargados")
            st.dataframe(build_catalog_summary(catalog_df), use_container_width=True, hide_index=True)

        global_summary = analysis["global"]
        trunks_df = analysis["trunks"].copy()
        indicators_df = analysis["indicators"].copy()
        strengths = rank_strengths(indicators_df, top_n=5)
        weaknesses = rank_weaknesses(indicators_df, top_n=5)

        st.subheader("Resumen ejecutivo del caso")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Perfil propuesto", analysis["profile"])
        m2.metric("Score global medio", f"{global_summary['score_raw_avg']:.2f} / 6")
        m3.metric("Vs objetivo global", format_pct(global_summary["vs_goal"]))
        m4.metric("Vs benchmark global", format_pct(global_summary["vs_bbdd"]))
        if analysis.get("profile_ratio") is not None:
            st.caption(f"Clasificación según Excel: {analysis['profile']} · ratio = {analysis['profile_ratio']:.2f} · ranking = {analysis.get('profile_rank', 'No disponible')}")
        st.markdown(build_executive_report(analysis, selected_species, selected_subspecies))

        if analysis.get("profile_thresholds"):
            threshold_df = pd.DataFrame([
                {"Ranking Excel": item["rank"], "Etiqueta": pretty_profile_label(item["label"]), "Corte (> )": item["cutoff"]}
                for item in sorted(analysis["profile_thresholds"], key=lambda x: x["rank"], reverse=True)
            ])
            st.dataframe(threshold_df, use_container_width=True, hide_index=True)

        st.subheader("Lectura por troncos")
        bar_df = trunks_df.copy()
        bar_df["tronco"] = bar_df["tronco"].apply(canonical_trunk_name)
        bar_df["Score técnico"] = bar_df["score_raw_avg"].round(2)
        bar_df["Referencia"] = 3.0
        bar_df["Benchmark aprox."] = (bar_df["score_raw_avg"] / bar_df["vs_bbdd"].replace({0: None})).fillna(0).round(2)
        bar_long = bar_df.melt(id_vars=["tronco"], value_vars=["Score técnico", "Referencia", "Benchmark aprox."], var_name="Serie", value_name="Valor")
        fig_bar = px.bar(bar_long, x="Valor", y="tronco", color="Serie", barmode="group", orientation="h", text="Valor", height=420)
        fig_bar.update_traces(texttemplate="%{text:.2f}", textposition="outside")
        fig_bar.update_layout(xaxis_title="Puntuación media", yaxis_title="Tronco", margin=dict(l=40, r=40, t=40, b=40))
        st.plotly_chart(fig_bar, use_container_width=True)

        st.subheader("Gráficos de araña para el informe")
        chart_cols = st.columns(3)
        for idx, (chart_title, chart_fig) in enumerate(report_charts):
            with chart_cols[idx % 3]:
                st.plotly_chart(chart_fig, use_container_width=True)
                st.caption(chart_title)

        left, right = st.columns(2)
        with left:
            st.markdown("### Top 5 fortalezas")
            strong_display = strengths[["tronco", "indicator", "score_raw", "vs_goal", "vs_bbdd"]].copy()
            strong_display["tronco"] = strong_display["tronco"].apply(canonical_trunk_name)
            strong_display["vs_goal"] = strong_display["vs_goal"].apply(format_pct)
            strong_display["vs_bbdd"] = strong_display["vs_bbdd"].apply(format_pct)
            strong_display.columns = ["Tronco", "Indicador", "Score", "Vs objetivo", "Vs benchmark"]
            st.dataframe(strong_display, use_container_width=True, hide_index=True)
        with right:
            st.markdown("### Top 5 debilidades")
            weak_display = weaknesses[["tronco", "indicator", "score_raw", "vs_goal", "vs_bbdd"]].copy()
            weak_display["tronco"] = weak_display["tronco"].apply(canonical_trunk_name)
            weak_display["vs_goal"] = weak_display["vs_goal"].apply(format_pct)
            weak_display["vs_bbdd"] = weak_display["vs_bbdd"].apply(format_pct)
            weak_display.columns = ["Tronco", "Indicador", "Score", "Vs objetivo", "Vs benchmark"]
            st.dataframe(weak_display, use_container_width=True, hide_index=True)

        st.subheader("Formación recomendada")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("### Áreas prioritarias de desarrollo")
            pr_display = priority_areas.copy()
            pr_display["tronco"] = pr_display["tronco"].apply(canonical_trunk_name)
            pr_display["priority"] = pr_display["priority"].round(2)
            pr_display.columns = ["Tronco", "Prioridad relativa"]
            st.dataframe(pr_display, use_container_width=True, hide_index=True)
            display_course_table("Formación base interna", base_internal)
            display_course_table("Formación base externa", base_external)
        with c2:
            st.markdown("### Áreas de especialización")
            sp_display = specialization_areas.copy()
            sp_display["tronco"] = sp_display["tronco"].apply(canonical_trunk_name)
            sp_display["priority"] = sp_display["priority"].round(2)
            sp_display.columns = ["Tronco", "Potencial relativo"]
            st.dataframe(sp_display, use_container_width=True, hide_index=True)
            display_course_table("Formación especializada interna", adv_internal)
            display_course_table("Formación especializada externa", adv_external)

        st.subheader("Preguntas del usuario para enriquecer el informe")
        with st.expander("Ver las 20 preguntas agrupadas", expanded=False):
            for group, questions in QUESTION_GROUPS.items():
                st.markdown(f"**{group}**")
                for q in questions:
                    st.write(f"- {q}")

        q_col1, q_col2 = st.columns([1, 2])
        with q_col1:
            selected_group = st.selectbox("Grupo de preguntas", list(QUESTION_GROUPS.keys()))
        with q_col2:
            selected_question = st.selectbox("Pregunta", QUESTION_GROUPS[selected_group])

        if st.button("Añadir pregunta al informe"):
            current_questions = [x["question"] for x in st.session_state["qa_items"]]
            if selected_question not in current_questions:
                st.session_state["qa_items"].append({
                    "group": selected_group,
                    "question": selected_question,
                    "answer": build_question_answer(selected_question, analysis),
                    "include": True,
                })
                st.rerun()

        if st.session_state["qa_items"]:
            for idx, item in enumerate(st.session_state["qa_items"]):
                with st.expander(f"{item['group']} — {item['question']}", expanded=False):
                    include = st.checkbox("Incorporar esta respuesta al informe final", value=item.get("include", True), key=f"include_{idx}")
                    answer = st.text_area("Respuesta editable", value=item.get("answer", ""), height=140, key=f"answer_{idx}")
                    if st.button("Eliminar esta pregunta", key=f"remove_{idx}"):
                        st.session_state["qa_items"].pop(idx)
                        st.rerun()
                    st.session_state["qa_items"][idx]["include"] = include
                    st.session_state["qa_items"][idx]["answer"] = answer

        st.subheader("Informe final editable")
        auto_report = build_final_report_text(analysis, selected_species, selected_subspecies, base_internal, base_external, adv_internal, adv_external, career_plan_df, st.session_state["qa_items"])
        if not st.session_state.get("final_report_text"):
            st.session_state["final_report_text"] = auto_report
        if st.button("Refrescar informe con el contenido actual"):
            st.session_state["final_report_text"] = build_final_report_text(analysis, selected_species, selected_subspecies, base_internal, base_external, adv_internal, adv_external, career_plan_df, st.session_state["qa_items"])
            st.rerun()

        final_report_text = st.text_area("Puedes editar libremente el informe antes de descargarlo", value=st.session_state["final_report_text"], height=520)
        st.session_state["final_report_text"] = final_report_text

        st.subheader("Acceso directo al buscador global")
        st.info("La pestaña ‘Buscador global de formación’ mantiene el buscador y su informe propios, separados del informe integrado de assessment y plan de carrera.")

        st.subheader("Descargas")
        excel_bytes = make_excel_export(analysis, base_internal, base_external, adv_internal, adv_external, career_plan_df, st.session_state["qa_items"])
        chart_images = [(title, figure_to_png_bytes(fig)) for title, fig in report_charts]
        docx_bytes = make_docx_bytes(final_report_text, chart_images=chart_images, asset_paths=asset_paths)
        pdf_bytes = make_pdf_bytes(final_report_text, chart_images=chart_images, asset_paths=asset_paths)
        html_bytes = make_html_report_bytes(final_report_text, analysis, selected_species, selected_subspecies, chart_images=chart_images, asset_paths=asset_paths)
        txt_bytes = final_report_text.encode("utf-8")

        d1, d2, d3, d4, d5 = st.columns(5)
        with d1:
            st.download_button("Descargar informe TXT", data=txt_bytes, file_name="informe_plan_carrera.txt", mime="text/plain", use_container_width=True)
        with d2:
            st.download_button("Descargar informe HTML", data=html_bytes, file_name="informe_plan_carrera.html", mime="text/html", use_container_width=True)
        with d3:
            st.download_button("Descargar informe DOCX", data=docx_bytes, file_name="informe_plan_carrera.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        with d4:
            st.download_button("Descargar informe PDF", data=pdf_bytes, file_name="informe_plan_carrera.pdf", mime="application/pdf", use_container_width=True)
        with d5:
            st.download_button("Descargar detalle Excel", data=excel_bytes, file_name="assessment_plan_carrera.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

with finder_tab:
    st.subheader("Buscador global de cursos, congresos y seminarios")
    st.write("Este módulo permite acceder directamente al buscador avanzado de oportunidades formativas en producción animal, con plantilla corporativa y descargas propias.")
    finder_html = load_finder_html(asset_paths)
    if finder_html:
        st.download_button(
            "Descargar buscador HTML independiente",
            data=finder_html.encode("utf-8"),
            file_name="buscador_formacion_produccion_animal_v2.html",
            mime="text/html",
            use_container_width=False,
        )
        components.html(finder_html, height=1950, scrolling=True)
    else:
        st.warning("No se ha localizado el archivo del buscador HTML. Añádelo al repositorio con el nombre 'buscador_formacion_produccion_animal_v2.html'.")
